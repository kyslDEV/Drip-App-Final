from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT_DOSSIE = DOCS / "DripTest_Dossie_Apresentacao_Producao.docx"
OUTPUT_ROTEIRO = DOCS / "DripTest_Roteiro_Executivo_Reuniao.docx"

TODAY = date(2026, 5, 25)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6773"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
GREEN = "E8F5E9"
YELLOW = "FFF7D6"
RED = "FCE8E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(table)


def set_borders(table, color="B8C0CC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)


def configure_document(title: str, subtitle: str, running_label: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = running_label
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(8)
    run = title_paragraph.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)

    meta = add_table(
        doc,
        ["Campo", "Informação"],
        [
            ["Projeto", "DripTest"],
            ["Data-base do documento", TODAY.strftime("%d/%m/%Y")],
            ["Base usada", "Arquivos do projeto, documentação existente, backend-python e schema PostgreSQL"],
        ],
        [2100, 7260],
    )
    for cell in meta.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph()
    return doc


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_borders(table)
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(INK)
        set_cell_shading(header_cells[index], LIGHT_GRAY)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
            set_cell_width(cells[index], widths_dxa[index])
    set_table_geometry(table, widths_dxa)
    return table


def add_callout(doc, title: str, body: str, fill: str = BLUE_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    set_borders(table, color="C8D3E0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_section_break(doc) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_dossie() -> None:
    doc = configure_document(
        "DripTest - Dossiê de Apresentação e Prontidão",
        "Funcionalidades atuais, estágio do projeto e caminho para uso em produção",
        "DripTest | Dossiê de apresentação e prontidão",
    )

    add_callout(
        doc,
        "Posicionamento recomendado",
        "O DripTest deve ser apresentado como uma ferramenta operacional funcional, madura para piloto controlado e validação com a área de qualidade. O fluxo principal já existe; a produção corporativa plena ainda depende de deploy oficial, governança, testes automatizados, auditoria completa e política de suporte.",
    )

    doc.add_heading("1. Resumo executivo", level=1)
    doc.add_paragraph(
        "O DripTest é um aplicativo para conduzir a análise de gotejamento do início ao fim: configuração operacional, pesagem inicial, cálculo do tempo previsto, acompanhamento da agenda, pesagem final, consolidação de resultados e geração de laudos."
    )
    add_bullets(
        doc,
        [
            "O frontend web/PWA já cobre o fluxo operacional principal e opera com armazenamento local.",
            "Há empacotamento Android via WebView, útil para campo e operação offline controlada.",
            "O backend FastAPI já existe no repositório, com PostgreSQL, autenticação básica, lotes, pesagens, sincronização e laudos oficiais.",
            "A integração está em transição: pesagem inicial e relatórios já conversam com a API, enquanto agenda e pesagem final ainda dependem majoritariamente do store local.",
            "Para produção corporativa, ainda faltam ambiente oficial, migrations, testes automatizados, política de backup, hardening de segurança, auditoria completa por alteração e assinatura/aprovação formal de laudos.",
        ],
    )

    doc.add_heading("2. Visão geral da solução", level=1)
    add_table(
        doc,
        ["Camada", "Estado atual", "Observação para apresentação"],
        [
            ["Web/PWA", "Funcional", "Pode ser demonstrado como fluxo operacional completo em navegador e modo instalável."],
            ["Offline local", "Funcional", "Dados permanecem no dispositivo por localStorage/store versionado, com uso independente da API."],
            ["Android WebView", "Empacotado", "Existe APK/debug e estrutura Android; produção exige assinatura, distribuição e homologação."],
            ["Backend FastAPI", "Implementado inicial", "API 0.2.0 cobre health, auth, lotes, pesagens, finalização, sync e laudos."],
            ["PostgreSQL", "Schema definido", "Modelo tem plants, users, lots, weighings, reports, sync_batches e audit_logs."],
            ["Governança", "Parcial", "Há base técnica; faltam processo formal, testes, backups, auditoria integral e operação assistida."],
        ],
        [1700, 2300, 5360],
    )

    doc.add_heading("3. Fluxo operacional coberto", level=1)
    add_numbered(
        doc,
        [
            "Configurar monitor, setor, lote e data de fabricação.",
            "Registrar pesagens iniciais com peso bruto, embalagem, peso líquido e tempo previsto.",
            "Calcular e acompanhar o cronograma de análise das amostras.",
            "Registrar a pesagem final e calcular perda/absorção absoluta e percentual.",
            "Consolidar resultados por lote, espécie, marca, monitor e setor.",
            "Gerar laudo técnico, exportar CSV, copiar texto, compartilhar e emitir PDF/impressão.",
            "Quando a API estiver configurada, sincronizar dados e emitir laudo oficial no backend.",
        ],
    )

    doc.add_heading("4. Funcionalidades por módulo", level=1)
    add_table(
        doc,
        ["Módulo", "Funcionalidades disponíveis", "Status"],
        [
            ["login.html", "Identifica monitor, setor, lote e fabricação; salva contexto operacional para reutilização.", "Pronto para identificação operacional; login corporativo fica em DripSettings/API."],
            ["DripTeste.html", "Registra pesagem inicial, calcula líquido inicial, tempo por peso bruto, interpolação e envia para POST /weighings quando API ativa.", "Pronto para piloto; integração com banco iniciada."],
            ["DripSchedule.html", "Monta agenda, calcula horários, status Devido/Próximo/Agendado e permite copiar resumo.", "Funcional local; migração para leitura central ainda pendente."],
            ["DripTestF.html", "Lista amostras, registra peso final, calcula líquido final, perda, percentual, indicador comercial, reabertura e resumo.", "Funcional local; endpoint de finalização já existe no backend, mas integração de tela ainda precisa avançar."],
            ["DripReports.html", "Consolida indicadores, sincroniza store, emite laudo oficial, consulta histórico, exporta CSV e gera PDF/impressão.", "Avançado para apresentação; precisa assinatura/aprovação formal para produção final."],
            ["DripSettings.html", "Configura URL/token da API, testa backend, faz login e sincronização manual.", "Base administrativa criada para transição local -> central."],
        ],
        [1650, 5450, 2260],
    )

    doc.add_heading("5. Dados, cálculos e regras de negócio", level=1)
    doc.add_paragraph("As principais regras já codificadas e documentadas incluem:")
    add_bullets(
        doc,
        [
            "Peso líquido inicial: peso bruto menos embalagem inicial.",
            "Tempo previsto pelo peso bruto com tabela de faixas e interpolação.",
            "Peso líquido final e perda/absorção absoluta.",
            "Perda/absorção percentual por amostra e média por lote.",
            "Indicador comercial por faixa percentual.",
            "Média de perda percentual final truncada em duas casas no padrão do laudo.",
            "Hash SHA-256 para pacote/laudo e número de laudo baseado em data/hash quando emitido.",
        ],
    )
    add_callout(
        doc,
        "Ponto de validação técnica",
        "A regra definitiva de perda/absorção deve ser formalmente aprovada pela área técnica antes de congelar banco, laudo oficial, auditoria e aplicativo móvel definitivo.",
        fill=YELLOW,
    )

    doc.add_heading("6. Laudos e rastreabilidade", level=1)
    doc.add_paragraph(
        "O módulo de laudos é o ponto mais forte para apresentação, porque mostra o valor do sistema: consolidação operacional transformada em documento técnico rastreável."
    )
    add_table(
        doc,
        ["Capacidade", "Estado atual"],
        [
            ["Prévia textual", "Gerada localmente a partir do store consolidado."],
            ["PDF/impressão", "Gerado no navegador com identificação, objetivo, método, resumo, tabelas e hash."],
            ["CSV", "Exportação e importação para backup, Excel e carga manual."],
            ["Laudo oficial no backend", "POST /reports salva snapshot JSON, hash SHA-256 e vínculos com pesagens."],
            ["Histórico oficial", "GET /reports e GET /reports/{report_id} consultam laudos emitidos."],
            ["Lacunas formais", "Assinatura, aprovação técnica, critério oficial de aceitação, versionamento e política de cancelamento/revisão."],
        ],
        [2600, 6760],
    )

    doc.add_heading("7. Arquitetura atual", level=1)
    add_table(
        doc,
        ["Componente", "Arquivos/estrutura", "Papel"],
        [
            ["Frontend", "DripTeste, DripSchedule, DripTestF, DripReports, DripSettings", "Interface operacional e geração local de relatórios."],
            ["Domínio JS", "drip-data.js", "Store, cálculos, normalização, relatórios, CSV e hash local."],
            ["Integração JS", "drip-api.js e drip-sync.js", "Configuração da API, sessão, requests e snapshot do store local."],
            ["PWA", "manifest.webmanifest e service-worker.js", "Instalação, cache e uso offline."],
            ["Backend", "backend-python/app", "API FastAPI com auth, lotes, pesagens, sync e laudos."],
            ["Banco", "database/schema.sql", "Modelo PostgreSQL com tabelas operacionais, laudos, sincronização e auditoria."],
            ["Android", "android-offline", "WebView com ativos web sincronizados para uso mobile offline."],
        ],
        [1800, 3450, 4110],
    )

    doc.add_heading("8. Matriz de prontidão", level=1)
    add_table(
        doc,
        ["Área", "Prontidão", "Leitura executiva"],
        [
            ["Fluxo operacional", "Alta para piloto", "O ciclo principal já pode ser demonstrado e validado em campo controlado."],
            ["Usabilidade", "Boa", "Telas têm navegação por etapa e resumo operacional; ajustes finos podem vir após piloto."],
            ["Dados locais", "Boa para piloto", "Funciona offline, mas produção requer backup e consolidação central."],
            ["Backend/banco", "Média", "Código implementado, mas precisa ambiente oficial, migrations, observabilidade e testes."],
            ["Segurança", "Média-baixa para produção", "Há token e sessão assinada no backend; a governança de usuários/perfis ainda precisa fechar."],
            ["Auditoria", "Parcial", "Tabela e logs existem, mas falta política completa para toda criação, alteração, reabertura, exclusão e aprovação."],
            ["Laudo oficial", "Média-alta", "Fluxo técnico existe; faltam assinatura, aprovação, critérios e controle documental formal."],
            ["Testes/QA", "Baixa para produção", "Não há suíte automatizada evidente; produção exige testes de domínio, API, integração e regressão."],
        ],
        [2300, 2300, 4760],
    )

    doc.add_heading("9. O que pode ser dito sobre uso e produção", level=1)
    add_table(
        doc,
        ["Pergunta", "Resposta recomendada"],
        [
            ["Já pode usar?", "Sim, para piloto controlado, validação operacional e demonstração com qualidade. Não como sistema corporativo final sem governança."],
            ["Já está em produção?", "Não deve ser posicionado como produção plena. O estado correto é MVP operacional com backend inicial e prontidão para piloto."],
            ["Os dados estão seguros?", "Para piloto, os dados locais e a sincronização atendem à validação. Produção requer banco oficial, backup, políticas de acesso e auditoria."],
            ["O laudo é formal?", "Ele é apresentável e tecnicamente estruturado. Para ser documento oficial corporativo, faltam assinatura, aprovação e controle documental."],
            ["O que falta para liberar?", "Homologar regra de negócio, implantar backend/banco, ativar login/perfis, automatizar testes, fechar auditoria e definir suporte."],
        ],
        [2400, 6960],
    )

    doc.add_heading("10. Lacunas e riscos principais", level=1)
    add_table(
        doc,
        ["Risco", "Impacto", "Ação recomendada"],
        [
            ["Dependência de localStorage", "Perda ou isolamento de dados por dispositivo.", "Usar piloto com exportação/backup e avançar sincronização central."],
            ["Regra de absorção não congelada formalmente", "Divergência entre tela, laudo e banco.", "Homologar regra com qualidade/processo antes de produção."],
            ["Ausência de testes automatizados evidentes", "Regressões em cálculo, laudo e API.", "Criar suíte mínima para domínio, frontend crítico e endpoints."],
            ["Migrations ainda pendentes", "Risco de alteração manual de banco.", "Introduzir Alembic ou processo formal de versionamento de schema."],
            ["Auditoria incompleta", "Dificuldade de comprovar histórico de alteração.", "Padronizar logs para criação, edição, finalização, reabertura, exclusão, emissão e cancelamento."],
            ["Laudo sem assinatura/aprovação", "Baixa força documental oficial.", "Adicionar responsável técnico, aprovador, versão, critérios e estado de aprovação."],
        ],
        [2500, 3150, 3710],
    )

    doc.add_heading("11. Plano recomendado para produção", level=1)
    add_table(
        doc,
        ["Fase", "Objetivo", "Entregáveis"],
        [
            ["1. Homologação funcional", "Validar fluxo e cálculo com qualidade.", "Casos de teste manuais, regra de absorção aprovada, layout final do laudo."],
            ["2. API e banco oficial", "Sair do uso puramente local.", "PostgreSQL em ambiente oficial, configuração segura, backups, migrations."],
            ["3. Integração completa das telas", "Reduzir divergência local/central.", "Finalização e reabertura via API, agenda lendo dados centrais/cache."],
            ["4. Segurança e auditoria", "Controlar acesso e rastreabilidade.", "Login/perfis no frontend, logs completos, política de edição/cancelamento."],
            ["5. QA e operação", "Preparar sustentação.", "Testes automatizados, checklist de release, monitoramento, plano de suporte."],
            ["6. Mobile", "Distribuir operação em campo.", "APK assinado/homologado ou evolução futura para Kotlin nativo com Room/WorkManager."],
        ],
        [1600, 3000, 4760],
    )

    doc.add_heading("12. Checklist mínimo antes de produção corporativa", level=1)
    add_bullets(
        doc,
        [
            "Ambiente de banco PostgreSQL oficial criado, com backups testados e política de retenção.",
            "Variáveis DRIP_DATABASE_URL, DRIP_API_TOKEN, CORS e bootstrap admin configuradas com segredos fortes.",
            "Migrations/versionamento de schema definido.",
            "Login real ativado no fluxo de uso, com perfis e permissões por papel.",
            "Regra de cálculo aprovada e coberta por testes automatizados.",
            "Fluxos de pesagem inicial, finalização, reabertura, sincronização e emissão de laudo testados ponta a ponta.",
            "Laudo com número, hash, responsável técnico, aprovador, assinatura/estado de aprovação e critério oficial.",
            "Auditoria completa para criação, alteração, finalização, reabertura, exclusão/cancelamento e emissão de laudo.",
            "Política de uso offline definida, incluindo o que fazer quando a sincronização falhar.",
            "APK ou PWA homologado nos dispositivos reais de operação.",
        ],
    )

    doc.add_heading("13. Materiais de apoio para apresentação", level=1)
    add_table(
        doc,
        ["Material", "Uso recomendado"],
        [
            ["Demonstração do fluxo", "Abrir login, pesagem inicial, cronograma, pesagem final e laudos nesta ordem."],
            ["Tela mais forte", "DripReports.html, porque evidencia consolidação, rastreabilidade, exportação e laudo."],
            ["Mensagem de fechamento", "O projeto já entrega valor operacional real e está pronto para validação; produção plena é o próximo ciclo."],
            ["Documentos existentes", "CAPACIDADES_ATUAIS, GESTAO_DADOS_LAUDOS_E_BANCO, INTEGRACAO_WEB_BANCO e backend README."],
        ],
        [2600, 6760],
    )

    doc.save(OUTPUT_DOSSIE)


def build_roteiro() -> None:
    doc = configure_document(
        "DripTest - Roteiro Executivo de Reunião",
        "Guia curto para apresentar funcionalidades, valor e estágio atual do projeto",
        "DripTest | Roteiro executivo",
    )

    add_callout(
        doc,
        "Mensagem principal",
        "O DripTest já é um MVP operacional: registra pesagens, calcula tempos e perdas, acompanha agenda, consolida dados e gera laudos. Ele está pronto para piloto controlado e validação com qualidade; ainda não deve ser tratado como produção corporativa plena.",
    )

    doc.add_heading("1. Abertura em 60 segundos", level=1)
    doc.add_paragraph(
        "O DripTest foi criado para padronizar a análise de gotejamento, reduzindo controles manuais e transformando registros de operação em dados consolidados e laudos rastreáveis. Hoje o fluxo principal já funciona em web/PWA e Android WebView, com backend FastAPI e PostgreSQL em fase inicial de integração."
    )

    doc.add_heading("2. Roteiro de 5 a 10 minutos", level=1)
    add_table(
        doc,
        ["Etapa", "O que mostrar", "Fala curta sugerida"],
        [
            ["1. Problema", "Processo manual, risco de erro e laudo disperso.", "O desafio é manter tempo, cálculo, lote e rastreabilidade consistentes do início ao fim."],
            ["2. Fluxo", "login -> pesagem inicial -> agenda -> final -> laudo.", "O app acompanha a operação por etapas e reduz retrabalho."],
            ["3. Funcionalidades", "Cálculos automáticos, status, exportação, PDF e sincronização.", "O projeto já entrega saídas concretas para a área de qualidade."],
            ["4. Arquitetura", "PWA, Android WebView, FastAPI e PostgreSQL.", "A base técnica já está preparada para sair do piloto local para uma operação centralizada."],
            ["5. Estágio", "Pronto para piloto, não para produção plena.", "O fluxo está maduro; a próxima fase é governança, segurança, auditoria, testes e deploy."],
            ["6. Próximo passo", "Homologação controlada.", "A decisão recomendada é validar em campo e fechar os requisitos de produção."],
        ],
        [1300, 3150, 4910],
    )

    doc.add_heading("3. Roteiro por slide", level=1)
    add_table(
        doc,
        ["Slide", "Título", "Conteúdo essencial"],
        [
            ["1", "DripTest", "Aplicativo para controle operacional da análise de gotejamento."],
            ["2", "Problema que resolve", "Registros dispersos, risco de erro manual, cronograma difícil e laudo pouco padronizado."],
            ["3", "Solução proposta", "Fluxo único para registrar, calcular, acompanhar, consolidar e gerar laudo."],
            ["4", "Fluxo operacional", "Monitor/lote -> pesagem inicial -> agenda -> pesagem final -> laudo."],
            ["5", "Funcionalidades atuais", "Cálculos, interpolação, status, resumo por lote, PDF, CSV, compartilhamento e backend inicial."],
            ["6", "Valor para qualidade", "Padronização, rastreabilidade, redução de erro, consolidação rápida e base para auditoria."],
            ["7", "Laudo técnico", "Número/hash, objetivo, método, rastreabilidade, resumo, tabelas e histórico oficial quando API ativa."],
            ["8", "Arquitetura", "Web/PWA, Android WebView, FastAPI, PostgreSQL e sincronização por snapshot."],
            ["9", "Estágio atual", "MVP operacional pronto para piloto controlado; produção plena ainda em preparação."],
            ["10", "Próximos passos", "Homologar regra, implantar banco/API, fechar auditoria, testes, login/perfis e assinatura do laudo."],
        ],
        [850, 2400, 6110],
    )

    doc.add_heading("4. Perguntas prováveis", level=1)
    add_table(
        doc,
        ["Pergunta", "Resposta curta"],
        [
            ["Já pode usar?", "Sim, para piloto controlado e validação operacional. Para produção plena, ainda faltam governança e hardening."],
            ["O backend já existe?", "Sim. O repositório já tem FastAPI, PostgreSQL, autenticação, lotes, pesagens, sincronização e laudos."],
            ["Então já está pronto para produção?", "Ainda não. Produção exige deploy oficial, backup, testes, auditoria, login/perfis e suporte."],
            ["O laudo já serve?", "Serve para apresentação e piloto; para uso oficial final, precisa assinatura, aprovação e critérios formais."],
            ["Qual é o maior risco?", "Congelar regras e operar sem testes/auditoria. A regra de perda/absorção deve ser homologada antes da produção."],
            ["Qual é o próximo passo?", "Piloto com casos reais, validação dos cálculos e fechamento do checklist de produção."],
        ],
        [2500, 6860],
    )

    doc.add_heading("5. Ordem recomendada da demonstração", level=1)
    add_numbered(
        doc,
        [
            "Abrir login.html e mostrar monitor, setor, lote e data de fabricação.",
            "Abrir DripTeste.html e cadastrar uma pesagem inicial com cálculo automático.",
            "Abrir DripSchedule.html e mostrar agenda/status das amostras.",
            "Abrir DripTestF.html e registrar peso final, perda e indicador.",
            "Abrir DripReports.html e mostrar consolidação, laudo, CSV/PDF e histórico oficial quando a API estiver ativa.",
            "Abrir DripSettings.html apenas se precisar demonstrar a transição para backend/API.",
        ],
    )

    doc.add_heading("6. Frase de fechamento", level=1)
    doc.add_paragraph(
        "O DripTest já resolve a operação principal e entrega valor real para padronização e rastreabilidade. A recomendação é tratar o sistema como pronto para piloto estruturado e usar o próximo ciclo para transformar essa base em produção corporativa com banco oficial, segurança, auditoria, testes e governança de laudos."
    )

    doc.save(OUTPUT_ROTEIRO)


if __name__ == "__main__":
    build_dossie()
    build_roteiro()
    print(OUTPUT_DOSSIE)
    print(OUTPUT_ROTEIRO)
