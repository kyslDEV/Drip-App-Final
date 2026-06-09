from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "docs" / "Formulario_Estudo_Caso_DripTest.md"
OUTPUT = ROOT / "docs" / "Formulario_Estudo_Caso_DripTest.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6773"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            if index < len(row.cells):
                set_cell_width(row.cells[index], width)
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_borders(table, color="B8C0CC", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr)
    run._r.append(field_end)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "DripTest | Formulario de estudo de caso"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    return doc


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Formulario - Estudo de Caso de Uso do DripTest")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "Documento de apoio para montar o formulario no Microsoft Forms, Google Forms ou ferramenta equivalente."
    )
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Objetivo", "Coletar evidencias de produtividade, eficiencia, qualidade percebida e impacto do DripTest."),
        ("Publico", "Usuarios finais, monitores, qualidade, producao, laboratorio, supervisao e gestao."),
        ("Tempo estimado", "5 a 8 minutos."),
        ("Uso das respostas", "Estruturar estudo de caso, indicadores antes x depois e melhorias futuras."),
    ]
    table.cell(0, 0).text = "Campo"
    table.cell(0, 1).text = "Descricao"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GRAY)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_geometry(table, [2300, 7060])
    set_borders(table)
    doc.add_paragraph()


def clean_text(text: str) -> str:
    text = text.strip()
    return text.replace("  ", " ")


def add_metadata_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if ":" in text:
        label, value = text.split(":", 1)
        r1 = p.add_run(label.strip() + ": ")
        r1.font.bold = True
        r1.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        p.add_run(value.strip())
    else:
        p.add_run(text)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, BLUE_GRAY)
    set_table_geometry(table, [9360])
    set_borders(table, color="C8D3E0")
    doc.add_paragraph()


def convert_markdown_body(doc: Document, text: str) -> None:
    in_options = False
    in_scale = False

    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped == "---":
            in_options = False
            in_scale = False
            continue

        if stripped.startswith("# "):
            in_options = False
            in_scale = False
            heading = stripped[2:].strip()
            if heading.startswith("Formulario -"):
                continue
            doc.add_heading(heading, level=1)
            continue

        if stripped.startswith("## "):
            in_options = False
            in_scale = False
            heading = stripped[3:].strip()
            level = 2
            if heading.startswith("Indicadores") or heading.startswith("Perguntas-chave"):
                level = 2
            doc.add_heading(heading, level=level)
            continue

        if stripped == "Opcoes:":
            in_options = True
            in_scale = False
            add_metadata_paragraph(doc, "Opcoes:")
            continue

        if stripped == "Escala:":
            in_options = False
            in_scale = True
            add_metadata_paragraph(doc, "Escala:")
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(clean_text(stripped[2:]), style="List Bullet")
            if in_options:
                paragraph.paragraph_format.space_after = Pt(3)
            continue

        if len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:4]:
            doc.add_paragraph(clean_text(stripped), style="List Number")
            continue

        if stripped.startswith("Tipo:") or stripped.startswith("Obrigatoria:") or stripped.startswith("Exemplo:"):
            add_metadata_paragraph(doc, stripped)
            continue

        if stripped.startswith("Use a escala") or stripped.startswith("As perguntas mais importantes"):
            add_callout(doc, stripped)
            continue

        if stripped.startswith("0 =") or stripped.startswith("10 ="):
            add_metadata_paragraph(doc, stripped)
            continue

        paragraph = doc.add_paragraph(clean_text(stripped))
        paragraph.paragraph_format.space_after = Pt(6)


def audit_docx(doc: Document) -> None:
    for section in doc.sections:
        assert round(section.page_width.inches, 2) == 8.5
        assert round(section.page_height.inches, 2) == 11.0
        assert round(section.left_margin.inches, 2) == 1.0
        assert round(section.right_margin.inches, 2) == 1.0
    assert len(doc.paragraphs) > 50
    assert len(doc.tables) >= 2


def main() -> None:
    text = INPUT.read_text(encoding="utf-8")
    doc = configure_document()
    add_title_block(doc)
    convert_markdown_body(doc, text)
    audit_docx(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
